from __future__ import annotations

import sys
from pathlib import Path


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from common.models import (  # noqa: E402
    CHUNKS_READY_EVENT_TYPE,
    DOC_METHOD,
    PAGE_COMPLETED_STATUS,
    PAGE_FAILED_STATUS,
    QUEUE_DOC,
    QUEUE_ENTITY,
    TEXT_EXTRACTION_COMPLETED_STATUS,
    TEXT_EXTRACTION_FAILED_STATUS,
    DocRoutedMessage,
    OutboxMessage,
    StoredFile,
)
from docs.extractor import extract_doc_document  # noqa: E402
from docs.worker import process_doc_payload, publish_pending_outbox, run_doc_worker  # noqa: E402


class FakeRepository:
    def __init__(self, stored_file: StoredFile):
        self.files = {stored_file.file_id: stored_file}
        self.saved_results = []
        self.outbox_by_key: dict[tuple[object, ...], OutboxMessage] = {}
        self.published = []
        self.errors = []

    def get_file(self, file_id: str) -> StoredFile | None:
        return self.files.get(file_id)

    def save_doc_result(self, result, publish_downstream: bool):
        entity_outbox_id = None
        if result.is_ready_for_entity and publish_downstream:
            key = ("entity", result.message.file_id)
            message = self.outbox_by_key.get(key)
            if message is None:
                message = OutboxMessage(
                    outbox_id="entity-1",
                    queue_name=QUEUE_ENTITY,
                    payload={
                        "event_type": CHUNKS_READY_EVENT_TYPE,
                        "file_id": result.message.file_id,
                        "source_queue_name": QUEUE_DOC,
                        "chunk_count": result.chunk_count,
                        "page_count": result.total_pages,
                    },
                )
                self.outbox_by_key[key] = message
            entity_outbox_id = message.outbox_id

        saved = result.with_outbox_ids(
            pages=result.pages,
            entity_outbox_id=entity_outbox_id,
        )
        self.saved_results.append(saved)
        return saved

    def list_pending_outbox(self, queue_name: str) -> list[OutboxMessage]:
        return [
            message
            for message in self.outbox_by_key.values()
            if message.queue_name == queue_name
        ]

    def mark_outbox_published(self, outbox_id: str) -> None:
        self.published.append(outbox_id)

    def record_outbox_error(self, outbox_id: str, error: str) -> None:
        self.errors.append((outbox_id, error))


class FakePublisher:
    def __init__(self):
        self.messages = []

    def publish(self, queue_name: str, payload: dict[str, object]) -> None:
        self.messages.append((queue_name, payload))


class FakeConsumer:
    def __init__(self, payloads: list[dict[str, object]]):
        self.payloads = payloads
        self.calls = []

    def consume(
        self,
        queue_name: str,
        handle_payload,
        max_messages: int | None = None,
        requeue_messages: bool = False,
    ) -> None:
        self.calls.append(
            {
                "queue_name": queue_name,
                "max_messages": max_messages,
                "requeue_messages": requeue_messages,
            }
        )
        payloads = self.payloads[:max_messages] if max_messages is not None else self.payloads
        for payload in payloads:
            handle_payload(payload)


class FakeMaterializer:
    def __init__(self, materialized_path: Path) -> None:
        self.materialized_path = materialized_path
        self.released = []

    def materialize(self, stored_file: StoredFile):
        class Materialized:
            def __init__(self, materialized_file: StoredFile) -> None:
                self.stored_file = materialized_file

        return Materialized(
            stored_file.with_materialized_path(str(self.materialized_path))
        )

    def release_if_final(self, stored_file: StoredFile, status: str | None) -> None:
        if status in {TEXT_EXTRACTION_COMPLETED_STATUS, TEXT_EXTRACTION_FAILED_STATUS}:
            self.released.append((stored_file.file_id, status))


def make_payload(**overrides: object) -> dict[str, object]:
    payload: dict[str, object] = {
        "schema_version": "2.0",
        "event_type": "file.routed",
        "run_id": "11111111-1111-1111-1111-111111111111",
        "file_id": "22222222-2222-2222-2222-222222222222",
        "routing_decision_id": "33333333-3333-3333-3333-333333333333",
        "source_type": "local",
        "source_uri": "local:///tmp/sample.txt",
        "external_id": "/tmp/sample.txt",
        "file_name": "sample.txt",
        "relative_path": "sample.txt",
        "extension": ".txt",
        "mime_type": "text/plain",
        "checksum_sha256": "a" * 64,
        "content_hash": None,
        "etag": None,
        "size_bytes": 123,
        "source_queue_name": "Queue-Archivos",
        "destination_queue_name": QUEUE_DOC,
        "route_type": "doc",
        "reason": "document_extension",
    }
    payload.update(overrides)
    return payload


def make_stored_file(path: Path, extension: str = ".txt") -> StoredFile:
    return StoredFile(
        file_id="22222222-2222-2222-2222-222222222222",
        run_id="11111111-1111-1111-1111-111111111111",
        source_type="local",
        source_uri=f"local://{path}",
        external_id=str(path),
        file_name=f"sample{extension}",
        relative_path=f"sample{extension}",
        extension=extension,
        mime_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            if extension == ".docx"
            else "text/plain"
            if extension == ".txt"
            else None
        ),
        size_bytes=path.stat().st_size if path.exists() else 0,
        checksum_sha256="a" * 64,
        content_hash=None,
        etag=None,
    )


def test_doc_routed_message_parses_router_v2_payload():
    message = DocRoutedMessage.from_payload(make_payload())

    assert message.schema_version == "2.0"
    assert message.destination_queue_name == QUEUE_DOC
    assert message.route_type == "doc"
    assert message.extension == ".txt"


def test_doc_routed_message_rejects_wrong_queue():
    payload = make_payload(destination_queue_name="Queue-PDF")

    try:
        DocRoutedMessage.from_payload(payload)
    except ValueError as exc:
        assert "Unsupported destination_queue_name" in str(exc)
    else:
        raise AssertionError("expected ValueError")


def test_doc_routed_message_rejects_wrong_route_type():
    payload = make_payload(route_type="pdf")

    try:
        DocRoutedMessage.from_payload(payload)
    except ValueError as exc:
        assert "Unsupported route_type" in str(exc)
    else:
        raise AssertionError("expected ValueError")


def test_doc_routed_message_rejects_unsupported_extension():
    payload = make_payload(extension=".pdf")

    try:
        DocRoutedMessage.from_payload(payload)
    except ValueError as exc:
        assert "Unsupported document extension" in str(exc)
    else:
        raise AssertionError("expected ValueError")


def test_extract_txt_utf8_writes_logical_page_and_chunks(tmp_path):
    path = tmp_path / "sample.txt"
    path.write_text("Nombre: Ana\n\nDireccion: Alameda 123", encoding="utf-8")
    message = DocRoutedMessage.from_payload(make_payload())

    result = extract_doc_document(message, make_stored_file(path))

    assert result.status == TEXT_EXTRACTION_COMPLETED_STATUS
    assert result.pages[0].method == DOC_METHOD
    assert result.pages[0].page_number == 1
    assert result.pages[0].status == PAGE_COMPLETED_STATUS
    assert result.pages[0].cpu_total_seconds is not None
    assert result.pages[0].peak_memory_mb is not None
    assert result.cpu_total_seconds >= 0
    assert result.chunk_count == 1
    assert "nombre: ana" in result.chunks[0].text
    assert result.chunks[0].source_map["segments"][0]["method"] == DOC_METHOD


def test_extract_txt_latin_1_fallback(tmp_path):
    path = tmp_path / "sample.txt"
    path.write_bytes("Cliente: José".encode("latin-1"))
    message = DocRoutedMessage.from_payload(make_payload())

    result = extract_doc_document(message, make_stored_file(path))

    assert result.status == TEXT_EXTRACTION_COMPLETED_STATUS
    assert "josé" in result.chunks[0].text
    assert result.chunks[0].source_map["segments"][0]["metadata"]["encoding"] == "latin-1"


def test_extract_docx_paragraphs_and_tables(tmp_path):
    from docx import Document

    path = tmp_path / "sample.docx"
    document = Document()
    document.add_paragraph("Cliente Ana")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Rut"
    table.cell(0, 1).text = "12.345.678-5"
    table.cell(1, 0).text = "Ciudad"
    table.cell(1, 1).text = "Santiago"
    document.save(path)
    payload = make_payload(
        file_name="sample.docx",
        relative_path="sample.docx",
        extension=".docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    message = DocRoutedMessage.from_payload(payload)

    result = extract_doc_document(message, make_stored_file(path, ".docx"))

    assert result.status == TEXT_EXTRACTION_COMPLETED_STATUS
    assert "cliente ana" in result.chunks[0].text
    assert "12.345.678-5" in result.chunks[0].text
    block_types = {
        segment["block_type"]
        for chunk in result.chunks
        for segment in chunk.source_map["segments"]
    }
    assert {"paragraph", "table"}.issubset(block_types)


def test_extract_empty_txt_completes_without_chunks(tmp_path):
    path = tmp_path / "empty.txt"
    path.write_text("", encoding="utf-8")
    message = DocRoutedMessage.from_payload(make_payload())

    result = extract_doc_document(message, make_stored_file(path))

    assert result.status == TEXT_EXTRACTION_COMPLETED_STATUS
    assert result.chunk_count == 0
    assert result.pages[0].reason == "empty_document"


def test_extract_corrupt_docx_fails_without_chunks(tmp_path):
    path = tmp_path / "sample.docx"
    path.write_bytes(b"not a docx")
    payload = make_payload(
        file_name="sample.docx",
        relative_path="sample.docx",
        extension=".docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    message = DocRoutedMessage.from_payload(payload)

    result = extract_doc_document(message, make_stored_file(path, ".docx"))

    assert result.status == TEXT_EXTRACTION_FAILED_STATUS
    assert result.pages[0].status == PAGE_FAILED_STATUS
    assert result.chunk_count == 0


def test_google_native_doc_fails_as_not_materialized(tmp_path):
    payload = make_payload(
        source_type="drive",
        source_uri="drive://file/doc-1",
        external_id="doc-1",
        file_name="drive-doc",
        relative_path="drive-doc",
        extension="",
        mime_type="application/vnd.google-apps.document",
        size_bytes=None,
        checksum_sha256=None,
        reason="google_document_mime",
    )
    message = DocRoutedMessage.from_payload(payload)
    stored_file = StoredFile(
        file_id=str(payload["file_id"]),
        run_id=str(payload["run_id"]),
        source_type="drive",
        source_uri="drive://file/doc-1",
        external_id="doc-1",
        file_name="drive-doc",
        relative_path="drive-doc",
        extension="",
        mime_type="application/vnd.google-apps.document",
        size_bytes=None,
        checksum_sha256=None,
        content_hash=None,
        etag="etag-1",
    )

    result = extract_doc_document(message, stored_file)

    assert result.status == TEXT_EXTRACTION_FAILED_STATUS
    assert result.pages[0].reason == "google_native_not_materialized"
    assert "prior export" in str(result.error)


def test_process_google_native_doc_uses_materialized_txt(tmp_path):
    exported_path = tmp_path / "drive-doc.txt"
    exported_path.write_text("Texto exportado desde Drive", encoding="utf-8")
    payload = make_payload(
        source_type="drive",
        source_uri="drive://file/doc-1",
        external_id="doc-1",
        file_name="Drive Doc",
        relative_path="Drive Doc",
        extension="",
        mime_type="application/vnd.google-apps.document",
        size_bytes=None,
        checksum_sha256=None,
        reason="google_document_mime",
    )
    stored_file = StoredFile(
        file_id=str(payload["file_id"]),
        run_id=str(payload["run_id"]),
        source_type="drive",
        source_uri="drive://file/doc-1",
        external_id="doc-1",
        file_name="Drive Doc",
        relative_path="Drive Doc",
        extension="",
        mime_type="application/vnd.google-apps.document",
        size_bytes=None,
        checksum_sha256=None,
        content_hash=None,
        etag="etag-1",
    )
    repository = FakeRepository(stored_file)
    materializer = FakeMaterializer(exported_path)

    result = process_doc_payload(payload, repository, materializer=materializer)

    assert result.status == TEXT_EXTRACTION_COMPLETED_STATUS
    assert result.chunk_count == 1
    assert result.chunks[0].text == "texto exportado desde drive"
    assert result.chunks[0].source_map["source_uri"] == "drive://file/doc-1"
    assert result.chunks[0].source_map["original_path"] == "drive://file/doc-1"
    assert materializer.released == [
        (str(payload["file_id"]), TEXT_EXTRACTION_COMPLETED_STATUS)
    ]


def test_process_doc_payload_publishes_entity_outbox(tmp_path):
    path = tmp_path / "sample.txt"
    path.write_text("Texto con datos", encoding="utf-8")
    repository = FakeRepository(make_stored_file(path))

    result = process_doc_payload(make_payload(), repository)

    assert result.status == TEXT_EXTRACTION_COMPLETED_STATUS
    assert result.entity_outbox_id == "entity-1"
    assert len(repository.list_pending_outbox(QUEUE_ENTITY)) == 1
    assert repository.list_pending_outbox(QUEUE_ENTITY)[0].payload["source_queue_name"] == QUEUE_DOC


def test_process_doc_payload_dev_mode_does_not_create_outbox(tmp_path):
    path = tmp_path / "sample.txt"
    path.write_text("Texto con datos", encoding="utf-8")
    repository = FakeRepository(make_stored_file(path))

    result = process_doc_payload(make_payload(), repository, publish_downstream=False)

    assert result.status == TEXT_EXTRACTION_COMPLETED_STATUS
    assert result.entity_outbox_id is None
    assert repository.list_pending_outbox(QUEUE_ENTITY) == []


def test_process_doc_payload_failure_does_not_create_outbox(tmp_path):
    path = tmp_path / "missing.txt"
    repository = FakeRepository(make_stored_file(path))

    result = process_doc_payload(make_payload(), repository)

    assert result.status == TEXT_EXTRACTION_FAILED_STATUS
    assert result.entity_outbox_id is None
    assert repository.list_pending_outbox(QUEUE_ENTITY) == []


def test_process_doc_payload_failure_releases_materialized_file(tmp_path):
    path = tmp_path / "missing.txt"
    repository = FakeRepository(make_stored_file(path))
    materializer = FakeMaterializer(path)

    result = process_doc_payload(
        make_payload(),
        repository,
        materializer=materializer,
    )

    assert result.status == TEXT_EXTRACTION_FAILED_STATUS
    assert result.entity_outbox_id is None
    assert materializer.released == [
        (str(make_payload()["file_id"]), TEXT_EXTRACTION_FAILED_STATUS)
    ]


def test_reprocess_doc_does_not_duplicate_entity_outbox(tmp_path):
    path = tmp_path / "sample.txt"
    path.write_text("Texto con datos", encoding="utf-8")
    repository = FakeRepository(make_stored_file(path))

    first = process_doc_payload(make_payload(), repository)
    second = process_doc_payload(make_payload(), repository)

    assert first.entity_outbox_id == "entity-1"
    assert second.entity_outbox_id == "entity-1"
    assert len(repository.list_pending_outbox(QUEUE_ENTITY)) == 1


def test_publish_pending_outbox_publishes_entity_messages(tmp_path):
    path = tmp_path / "sample.txt"
    path.write_text("Texto con datos", encoding="utf-8")
    repository = FakeRepository(make_stored_file(path))
    process_doc_payload(make_payload(), repository)
    publisher = FakePublisher()

    published = publish_pending_outbox(repository, publisher)

    assert published == 1
    assert publisher.messages[0][0] == QUEUE_ENTITY
    assert repository.published == ["entity-1"]


def test_run_doc_worker_uses_source_queue_and_requeue_flag(tmp_path):
    path = tmp_path / "sample.txt"
    path.write_text("Texto con datos", encoding="utf-8")
    repository = FakeRepository(make_stored_file(path))
    publisher = FakePublisher()
    consumer = FakeConsumer([make_payload()])

    run_doc_worker(
        repository=repository,
        publisher=publisher,
        consumer=consumer,
        publish_downstream=False,
        max_messages=1,
        requeue_messages=True,
    )

    assert consumer.calls == [
        {
            "queue_name": QUEUE_DOC,
            "max_messages": 1,
            "requeue_messages": True,
        }
    ]
    assert repository.saved_results[0].status == TEXT_EXTRACTION_COMPLETED_STATUS
